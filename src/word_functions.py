from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.run import Run

import settings

def convert_to_rgb(color_hex)->RGBColor:
    """
    :param color_hex: string: #RRGGBB

    return RGBColor() instance
    """
    r = int(color_hex[1:3], 16)
    g = int(color_hex[3:5], 16)
    b = int(color_hex[5:7], 16)

    return RGBColor(r, g, b)

def shade_cell(cell: _Cell, color_str: str):
    """
    Word XML magic
    :param cell: the _Cell object that will get shaded
    :param color_str: a string representation of the rgb color, can be with or without #
    """
    cell_element = cell._tc

    shading_elm = OxmlElement("w:shd")

    shading_elm.set(qn("w:fill"), color_str)

    cell_element.get_or_add_tcPr().append(shading_elm)

def set_vertical_alignment(cell: _Cell, align:str="center"):
    """
    Sets the vertical alignment of a cell, default is center aligned
    :param cell: the _Cell object of which alignment will be set
    :param align: can be top, center og bottom
    """

    if align not in('top', 'center', 'bottom'):
        raise ValueError("Invalid alignment value.")
    
    cell_element = cell._tc

    cell_properties = cell_element.get_or_add_tcPr()

    vAlign_elm = OxmlElement('w:vAlign')
    vAlign_elm.set(qn('w:val'), align)

    cell_properties.append(vAlign_elm)

def set_horizontal_alignment(cell:_Cell, align:WD_ALIGN_PARAGRAPH):
    """
    Set horizontal alignment of text in a table cell.

    :param cell: a cell in a Word table
    :param align: a member of the WD_ALIGN_PARAGRAPH enumeration
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
    
def set_text_color(run, color):
    """
    Sets the color of the text in a run
    :param run: a text run of a paragraph
    :param color: a color string
    """

    run.font.color.rgb = convert_to_rgb(color)

def insert_text_in_cell(cell:_Cell, text:str, alignment:WD_ALIGN_PARAGRAPH=None, font:str="Calibri", size=11, bold:bool=False)->Run:
    """
    Inserts text in a cell, returns the run containing the text
    :param cell: the cell where the text will be inserted
    :param text: the text to insert
    :param alignment: optioinal, must be a WD_ALIGN_PARAGRAPH type
    :param font: optional, defaults to Calibri
    :param size: optional, defaults to 11
    :param bold: optional
    """
    if alignment != None:
        set_horizontal_alignment(cell, alignment)

    run = cell.paragraphs[0].add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold

    return run

def create_table(doc:Document, postname:str, rotation:list):
    """
    Creates the table for scoring and order of groups
    :param doc: the document in which the table will be established
    :param postname: the name of the post
    :param rotation: ordered list of each 
    """
    table = doc.add_table(rows=len(rotation)+1, cols=4)
    legend = ["Klasse 1", "Poeng", "Klasse 2", "Poeng"]

    for i, cell in enumerate(table.rows[0].cells):
        txt_run = insert_text_in_cell(cell, legend[i], WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
        shade_cell(cell, settings.COLORS["dark"])
        set_text_color(txt_run, settings.COLORS["light"])
    
    for group, row in zip(rotation, list(table.rows)[1:]):
        cell1 = row.cells[0]
        cell2 = row.cells[2]
        insert_text_in_cell(cell1, group[0])
        insert_text_in_cell(cell2, group[1])

    table.style = "Table Grid"

